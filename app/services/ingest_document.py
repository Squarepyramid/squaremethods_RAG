"""
SquareMethods - Document Ingestion Service
==========================================
Parses uploaded equipment manuals (PDF or Word) and stores
chunked embeddings in knowledge_embeddings. Also (see IMAGE_EXTRACTION_ENABLED
below) can extract embedded images from PDF manuals and store them in S3 +
equipment_manual_images, so PM strategy generation can reference the
manual's OWN component images instead of a generic stock photo pulled
from the open web.

*** IMAGE EXTRACTION IS CURRENTLY DISABLED (IMAGE_EXTRACTION_ENABLED = False) ***
Per-image S3 uploads were adding meaningful latency to ingestion,
especially on manuals with many embedded drawings. Turned off for now so
ingestion stays fast; all the image extraction/storage code below is
intact and unused, not deleted -- flip the flag back to True to resume.
See that constant's comment for what changes when you do. Nothing about
re-enabling it requires re-ingesting already-processed documents' TEXT
chunks; only documents ingested while the flag is off will be missing
images until they're re-ingested.

*** TEXTRACT-ONLY BUILD -- COMPARISON VARIANT, NOT THE RECOMMENDED DEFAULT ***
This file is a variant of document_ingest_service.py that sends EVERY
PDF through Amazon Textract, unconditionally -- no pypdf native-text
pass, no scanned-vs-clean detection, no per-page or per-document
threshold. It exists so the two approaches can be compared side by
side (behavior, latency, cost, log output) on the same real documents.
It is NOT what's currently recommended for production: for a mix of
scanned and normal digitally-generated manuals, this build pays
Textract's async-job latency (and per-page cost) on every single
ingest, including the normal manuals that already have a perfectly
good text layer and previously extracted for free in milliseconds via
pypdf. See document_ingest_service.py (the hybrid version) for the
native-first, Textract-only-when-needed design and the reasoning
behind it.

We tried a local-OCR fix first (pymupdf to render pages + pytesseract +
a `tesseract` binary in the Lambda image) and deliberately backed out of
it in favor of Amazon Textract. Not because local OCR can't work in
principle, but because of what actually happened trying to deploy it on
this Lambda's base image (public.ecr.aws/lambda/python:3.11, Amazon
Linux 2): AL2's EPEL7 repo -- the only yum source for `tesseract` -- is
frozen at a decade-old 3.04.00 build (EL7 went EOL mid-2024) that
predates Tesseract's modern LSTM engine and that pytesseract can't even
reliably parse the version string of; AL2023 doesn't package tesseract
at all; and the version-string failure took down the whole Lambda
invocation via an uncaught SystemExit (pytesseract raises SystemExit,
not Exception, when it can't parse `tesseract --version` output) before
we even got to figure out the OCR quality question. Getting a real
tesseract binary onto this image was going to mean either compiling it
from source against AL2's old glibc/gcc (fragile, slow CI builds, needs
a modern-enough C++ toolchain that AL2 also doesn't ship) or pulling
prebuilt binaries from a third party layer project.

Textract sidesteps all of that: it's a plain boto3 API call (`boto3` was
already a dependency), so there is no OS package, no binary, no Docker
build step, and nothing that can break in CI for glibc/version reasons.
See TEXTRACT_ENABLED below.

How it's used here: Textract's async document-text-detection API only
takes input from S3 (not raw bytes), and it always processes the ENTIRE
submitted document -- there's no way to ask it to OCR just specific
pages, and there's no billing difference between submitting the whole
file vs. a subset (you're charged per page of whatever you submit
either way). In THIS build, every PDF is submitted to Textract exactly
once, unconditionally -- there is no native-extraction check and no
fallback if the Textract call fails (contrast with the hybrid version,
which falls back to native text on a Textract error since it already
has that text in hand; this build has nothing to fall back to, so a
Textract failure here fails the whole ingest job).

DEPLOYMENT REQUIREMENT for Textract: the Lambda's execution role needs
    textract:StartDocumentTextDetection
    textract:GetDocumentTextDetection
    textract:StartDocumentAnalysis   -- added for the LAYOUT/Markdown extraction path
    textract:GetDocumentAnalysis     -- (see extract_markdown_from_pdf()); the plain
                                         text-detection actions above are still needed
                                         too, for the unchanged extract_text_from_pdf()
                                         path (image context_text, still used if
                                         IMAGE_EXTRACTION_ENABLED is turned back on)
    s3:GetObject on the bucket manuals are uploaded to
file_url is expected to already be an S3 URL (that's how every document
reaches _run_ingest today) and is used directly as Textract's input --
no re-upload needed in the common case. If a future caller passes a
non-S3 URL, the already-downloaded bytes are uploaded to a scratch S3
key under S3_BUCKET for the Textract call and deleted afterward.

Key design decisions:
  - Each chunk gets a stable UUID derived from (doc_uuid, chunk_index)
    via uuid5, so re-ingesting the same file produces the same chunk UUIDs
  - doc_uuid is embedded in the content prefix so clear_document_chunks
    can delete all chunks for a specific document without a separate column
  - Deleting a document removes only its chunks, not other documents
  - Deleting a node removes all chunks for that node
  - Ingest runs as an async background job via SQS (see run_ingest_job)
  - Embeddings are generated concurrently in bounded batches to handle
    large documents (50+ pages) within the Lambda timeout window
  - EVERY PDF's text extraction goes through Amazon Textract, unconditionally
    (this is the Textract-only comparison build -- see the module docstring
    section above). No pypdf native-text pass, no scanned-vs-clean check.
  - Extracted manual images are associated with their PAGE's text as
    "context_text" (not a precise per-figure caption) -- OEM manuals
    rarely have clean one-image-per-component layouts; a parts-list
    page usually has one exploded-view image covering many components,
    with a table of callouts below it. Page-level association is a
    reasonable middle ground: coarser than a true caption match, but
    still specific to the actual equipment, not a generic web photo.
    That page text now goes through the same native-or-Textract
    extraction as the main text pass, so scanned pages get real
    context_text too once IMAGE_EXTRACTION_ENABLED is turned back on.
  - Extracted images are size-filtered (MIN_IMAGE_BYTES) as a first-pass
    heuristic to skip small logos/watermarks/icons -- not perfect, but
    avoids flooding storage with dozens of duplicate section-header
    logos that OEM manuals commonly repeat on every divider page.
  - S3 stores images privately; only the S3 KEY is saved in the DB, not
    a permanent public URL. A presigned URL is generated on demand
    whenever an image is actually referenced (e.g. during PM strategy
    generation), since these manuals may be confidential OEM documents
    and shouldn't be permanently public on the open internet.
  - Each text chunk is additionally scanned (regex, no model call) for
    OEM-style parts-list rows ("Lv Item Component-no Description Qty
    Un", e.g. "2 0001 341047 02 RAIL-TOP-RH 1.000 EA") and any
    (item_no, part_no) pairs found are captured in the chunk's content
    prefix as "bom_items:...". This costs nothing at ingest time and
    is a no-op (empty field) for manuals that don't use this table
    format. It isn't consumed by anything yet -- it's stored so a
    future item-number-precise image-linking feature doesn't require
    re-ingesting every manual to backfill it. See extract_bom_items()
    docstring for the matching caveat. Note this regex is matched
    against extracted text, so on Textract'd pages its hit rate depends
    on Textract's read quality on that particular scan -- a bad match
    is a safe no-op, not an error. (Textract also has a native table-
    extraction mode that would read these rows as structured cells
    instead of guessing at row shapes with a regex -- worth considering
    for a future pass, not part of this change.)

REQUIRES A DB MIGRATION before this will work -- see the
equipment_manual_images table DDL in the module docstring below the
imports. This has NOT been run automatically; it must be applied to
your database before deploying this file.

-- Migration required:
-- CREATE TABLE equipment_manual_images (
--     id UUID PRIMARY KEY,
--     company_id UUID NOT NULL REFERENCES companies(id),
--     equipment_id UUID NOT NULL REFERENCES equipment(id) ON DELETE CASCADE,
--     doc_id UUID NOT NULL,
--     s3_key TEXT NOT NULL,
--     page_number INT,
--     context_text TEXT,
--     created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
--     updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
-- );
-- CREATE INDEX equipment_manual_images_equipment_id ON equipment_manual_images(equipment_id);
-- CREATE INDEX equipment_manual_images_company_id ON equipment_manual_images(company_id);
"""

import io
import os
import re
import time
import uuid
import hashlib
import logging
import asyncio
import urllib.parse
import requests

import boto3
import psycopg2.extras

from app.utils.db import get_db_connection
from app.services.embeddings import get_embedding

log = logging.getLogger(__name__)

CHUNK_SIZE         = 500   # approximate words per chunk
CHUNK_OVERLAP      = 50    # words overlap between chunks
EMBED_CONCURRENCY  = 10    # max simultaneous embedding calls in flight

S3_BUCKET  = os.environ.get("S3_BUCKET", "squaremethods")
AWS_REGION = os.environ.get("AWS_REGION", "ca-central-1")

# Master switch for the PDF image-extraction + S3-upload step. Off for now
# to keep ingestion fast -- extracting every embedded image and uploading
# each one to S3 synchronously was the slow part of ingest, not the text
# chunking/embedding. When you're ready to bring images back:
#   1. Flip this to True.
#   2. That's it -- extract_images_from_pdf(), save_manual_images(), and
#      clear_document_images() are all still here and unchanged, just not
#      being called from _run_ingest() while this is False.
# If ingestion needs to stay fast even with images back on, the next step
# would be moving the S3 uploads to their own async/background step rather
# than inline in the main ingest job -- not needed until this flag is on.
IMAGE_EXTRACTION_ENABLED = False

# Heuristic filter for skipping small logos/icons/watermarks extracted
# from a PDF. Tune based on what your actual manuals contain -- OEM
# section-divider logos in the manuals we've tested were consistently
# well under this size, while real component photos/diagrams were not.
MIN_IMAGE_BYTES = 8000

_EXTENSION_CONTENT_TYPES = {
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif":  "image/gif",
    ".bmp":  "image/bmp",
    ".tiff": "image/tiff",
    ".webp": "image/webp",
}


# ── File URL to stable UUID ───────────────────────────────────────────────────

def url_to_uuid(file_url: str) -> str:
    """
    Generate a stable UUID from the file URL.
    Same URL always produces the same UUID so we can reliably
    identify and delete chunks for a specific document.
    """
    return str(uuid.UUID(hashlib.md5(file_url.encode()).hexdigest()))


# ── Text extraction: Amazon Textract, unconditionally ─────────────────────────
#
# TEXTRACT-ONLY BUILD: no native-text check, no per-page or per-document
# threshold. Every PDF is submitted to Textract exactly once. See the
# module docstring for why this exists as a separate build from the
# hybrid (native-first) version.

TEXTRACT_ENABLED               = True   # requires the IAM permissions in the module docstring
TEXTRACT_POLL_INTERVAL_SECONDS = 3      # how often to check job status while waiting
TEXTRACT_MAX_WAIT_SECONDS      = 300    # give up waiting after this long (leaves headroom under the Lambda timeout for download/chunk/embed)

_S3_VIRTUAL_HOSTED_RE = re.compile(r"^(?P<bucket>[^.]+)\.s3(?:[.-][a-z0-9-]+)?\.amazonaws\.com$")
_S3_PATH_STYLE_RE     = re.compile(r"^s3(?:[.-][a-z0-9-]+)?\.amazonaws\.com$")


def _parse_s3_url(file_url: str):
    """
    Best-effort parse of an S3 HTTPS URL into (bucket, key). Handles both
    virtual-hosted-style (https://bucket.s3.amazonaws.com/key) and
    path-style (https://s3.amazonaws.com/bucket/key) URLs, with or
    without a region segment, and strips any query string (e.g. from a
    presigned URL) before treating the rest as the key.

    Returns None if file_url doesn't look like an S3 URL -- callers fall
    back to uploading the already-downloaded bytes to S3_BUCKET instead
    of failing outright.
    """
    try:
        parsed = urllib.parse.urlparse(file_url)
        host   = parsed.netloc
        path   = parsed.path.lstrip("/")
        if not host or not path:
            return None

        m = _S3_VIRTUAL_HOSTED_RE.match(host)
        if m:
            return m.group("bucket"), urllib.parse.unquote(path)

        if _S3_PATH_STYLE_RE.match(host):
            bucket, _, key = path.partition("/")
            if bucket and key:
                return bucket, urllib.parse.unquote(key)

        return None
    except Exception:
        return None


def _textract_client():
    return boto3.client("textract", region_name=AWS_REGION)


def _textract_extract_document(file_bytes: bytes, file_url: str) -> dict:
    """
    Run the whole PDF through Textract's async document-text-detection
    API and return {page_number (1-indexed): text}.

    Textract's async API only accepts input from S3, and it always
    processes the entire submitted document -- there's no per-page
    dispatch and no cost difference between submitting the whole file
    vs. a subset (billing is per page of whatever's submitted either
    way). So unlike a local-OCR approach (where CPU time made
    per-page selectivity worth the complexity), this always submits the
    full document once ANY page needs it, and uses Textract's text for
    every page it covers -- simpler, and Textract's LINE-level output is
    generally at least as clean as pypdf's raw text for pages that
    already had a text layer too.

    file_url is expected to already point at an S3 object -- that's how
    every document reaches _run_ingest today (see the S3 URL logged in
    _run_ingest). If it doesn't parse as one, the already-downloaded
    bytes are uploaded to a scratch key under S3_BUCKET instead, used
    for the Textract call, then deleted.
    """
    textract = _textract_client()

    parsed = _parse_s3_url(file_url)
    scratch_bucket = scratch_key = None
    if parsed:
        bucket, key = parsed
    else:
        log.warning(
            f"file_url doesn't look like an S3 URL ({file_url!r}); "
            f"uploading a scratch copy to S3_BUCKET for Textract instead."
        )
        scratch_bucket = bucket = S3_BUCKET
        scratch_key    = key    = f"textract-scratch/{uuid.uuid4()}.pdf"
        _s3_client().put_object(Bucket=bucket, Key=key, Body=file_bytes)

    try:
        start_response = textract.start_document_text_detection(
            DocumentLocation={"S3Object": {"Bucket": bucket, "Name": key}}
        )
        job_id = start_response["JobId"]

        waited = 0
        result = None
        while True:
            result = textract.get_document_text_detection(JobId=job_id)
            status = result["JobStatus"]
            if status in ("SUCCEEDED", "PARTIAL_SUCCESS"):
                break
            if status == "FAILED":
                raise RuntimeError(
                    f"Textract job {job_id} failed: {result.get('StatusMessage')}"
                )
            if waited >= TEXTRACT_MAX_WAIT_SECONDS:
                raise TimeoutError(
                    f"Textract job {job_id} still {status} after "
                    f"{TEXTRACT_MAX_WAIT_SECONDS}s -- giving up for this ingest."
                )
            time.sleep(TEXTRACT_POLL_INTERVAL_SECONDS)
            waited += TEXTRACT_POLL_INTERVAL_SECONDS

        # Collect every LINE block across all result pages. Textract
        # paginates the Blocks array itself via NextToken -- unrelated to
        # the document's own page numbers, which come from each block's
        # "Page" field.
        pages_lines = {}
        next_token  = None
        while True:
            if next_token:
                result = textract.get_document_text_detection(
                    JobId=job_id, NextToken=next_token
                )
            for block in result.get("Blocks", []):
                if block.get("BlockType") == "LINE" and block.get("Text"):
                    page_num = block.get("Page", 1)
                    pages_lines.setdefault(page_num, []).append(block["Text"])
            next_token = result.get("NextToken")
            if not next_token:
                break

        return {page_num: "\n".join(lines) for page_num, lines in pages_lines.items()}
    finally:
        if scratch_key:
            try:
                _s3_client().delete_object(Bucket=scratch_bucket, Key=scratch_key)
            except Exception as e:
                log.warning(f"Failed to delete Textract scratch object {scratch_key}: {e}")


# ── Text extraction: Textract Layout Analysis -> Markdown ─────────────────────
#
# Added for structural (not flat-word-count) chunking that generalizes
# across manuals with different formatting conventions -- see the earlier
# ALL-CAPS-regex header detector this replaces, which was reverse-
# engineered from ONE manual's casing style and had no reason to
# generalize to a different manufacturer's Title Case or numbered
# headers. Textract's LAYOUT feature (AnalyzeDocument / its async
# StartDocumentAnalysis form) classifies each block of a page as a real
# structural element -- LAYOUT_TITLE, LAYOUT_SECTION_HEADER, LAYOUT_TEXT,
# LAYOUT_LIST, LAYOUT_TABLE, etc. -- from the document's actual visual
# layout (font size, position, spacing), not any assumption about a
# particular manual's text formatting. Still just a boto3 API call, same
# "no OS package, nothing to break in CI" property that got Textract
# chosen over local OCR in the first place (see the module docstring) --
# this doesn't reopen that already-settled tradeoff.
#
# Kept SEPARATE from _textract_extract_document()/_extract_pdf_page_texts()
# above rather than replacing them -- those still back
# extract_images_from_pdf()'s page-level context_text (out of scope here;
# IMAGE_EXTRACTION_ENABLED is currently False anyway) and anything else
# that specifically wants flat per-page text with no structure.
#
# *** UNVERIFIED AGAINST A LIVE CALL *** -- this sandbox has no AWS
# credentials. The block-shape assumptions below (CHILD relationships
# pointing to LINE blocks; which BlockTypes appear under FeatureTypes=
# ["LAYOUT"]) come from Textract's documented API, not a real response
# inspected here. Smoke-test against one real manual before trusting this
# in production -- log the raw Blocks for that run if the resulting
# Markdown looks wrong (missing sections, headers not detected, garbled
# ordering) so there's something concrete to debug against.

_LAYOUT_HEADER_TYPES = {"LAYOUT_TITLE": "#", "LAYOUT_SECTION_HEADER": "##"}
_LAYOUT_SKIP_TYPES    = {"LAYOUT_HEADER", "LAYOUT_FOOTER", "LAYOUT_PAGE_NUMBER", "LAYOUT_FIGURE"}
_PAGE_MARKER_RE = re.compile(r"<!--\s*page:(\d+)\s*-->")


def _textract_analyze_layout(file_bytes: bytes, file_url: str) -> list:
    """
    Same async-job/S3-input/polling/pagination pattern as
    _textract_extract_document() above, but calls StartDocumentAnalysis
    with FeatureTypes=["LAYOUT"] instead of StartDocumentTextDetection,
    and returns the FULL raw Blocks list (LAYOUT_* blocks AND the
    underlying LINE/WORD blocks they reference via CHILD relationships)
    rather than a pre-reduced page->text dict, since _blocks_to_markdown()
    below needs to resolve each LAYOUT block's children itself.
    """
    textract = _textract_client()

    parsed = _parse_s3_url(file_url)
    scratch_bucket = scratch_key = None
    if parsed:
        bucket, key = parsed
    else:
        log.warning(
            f"file_url doesn't look like an S3 URL ({file_url!r}); "
            f"uploading a scratch copy to S3_BUCKET for Textract instead."
        )
        scratch_bucket = bucket = S3_BUCKET
        scratch_key    = key    = f"textract-scratch/{uuid.uuid4()}.pdf"
        _s3_client().put_object(Bucket=bucket, Key=key, Body=file_bytes)

    try:
        start_response = textract.start_document_analysis(
            DocumentLocation={"S3Object": {"Bucket": bucket, "Name": key}},
            FeatureTypes=["LAYOUT"],
        )
        job_id = start_response["JobId"]

        waited = 0
        result = None
        while True:
            result = textract.get_document_analysis(JobId=job_id)
            status = result["JobStatus"]
            if status in ("SUCCEEDED", "PARTIAL_SUCCESS"):
                break
            if status == "FAILED":
                raise RuntimeError(
                    f"Textract layout-analysis job {job_id} failed: {result.get('StatusMessage')}"
                )
            if waited >= TEXTRACT_MAX_WAIT_SECONDS:
                raise TimeoutError(
                    f"Textract layout-analysis job {job_id} still {status} after "
                    f"{TEXTRACT_MAX_WAIT_SECONDS}s -- giving up for this ingest."
                )
            time.sleep(TEXTRACT_POLL_INTERVAL_SECONDS)
            waited += TEXTRACT_POLL_INTERVAL_SECONDS

        blocks = []
        next_token = None
        while True:
            if next_token:
                result = textract.get_document_analysis(JobId=job_id, NextToken=next_token)
            blocks.extend(result.get("Blocks", []))
            next_token = result.get("NextToken")
            if not next_token:
                break

        return blocks
    finally:
        if scratch_key:
            try:
                _s3_client().delete_object(Bucket=scratch_bucket, Key=scratch_key)
            except Exception as e:
                log.warning(f"Failed to delete Textract scratch object {scratch_key}: {e}")


def _resolve_block_text(block: dict, blocks_by_id: dict, join: str = " ") -> str:
    """
    A LAYOUT_* block doesn't carry its own text directly -- it references
    the LINE blocks that make it up via a CHILD relationship. Joins their
    Text fields in the order given (Textract returns child ids in reading
    order per its docs). Defensive: a missing/unresolvable child is
    skipped, not fatal -- a chunk with a small gap in it is far better
    than crashing the whole ingest job on a block shape this sandbox
    couldn't verify live.
    """
    pieces = []
    for rel in block.get("Relationships", []) or []:
        if rel.get("Type") != "CHILD":
            continue
        for child_id in rel.get("Ids", []):
            child = blocks_by_id.get(child_id)
            if child and child.get("Text"):
                pieces.append(child["Text"])
    return join.join(pieces).strip()


def _blocks_to_markdown(blocks: list) -> str:
    """
    Walks Textract's LAYOUT blocks (see _textract_analyze_layout()) in
    reading order and synthesizes one Markdown string for the whole
    document, with an invisible `<!-- page:N -->` HTML-comment marker
    inserted at every page boundary. Confirmed (mechanically tested
    against langchain_text_splitters directly, alongside this change)
    that MarkdownHeaderTextSplitter preserves these markers inside
    whatever chunk's content they fall into -- that's what lets
    chunk_markdown() recover a page_start/page_end per chunk after
    splitting, then strip the markers back out before the text is
    embedded or shown to anyone.

    LAYOUT_HEADER/LAYOUT_FOOTER/LAYOUT_FIGURE are skipped from chunk
    CONTENT -- running page headers/footers aren't useful body text, and
    figures are handled by the separate (currently disabled) image-
    extraction path, not text content. LAYOUT_TABLE is emitted as plain
    paragraph text rather than a real Markdown table for now -- table
    cell geometry parsing is a real gap, flagged rather than silently
    done wrong; a garbled-but-present table beats a missing one, but
    revisit if a manual with data-critical tables (specs, tolerances)
    shows extraction quality problems here.

    LAYOUT_PAGE_NUMBER blocks are excluded from content the same way, but
    NOT thrown away -- their text is the manual's own PRINTED page label,
    which is what a real page marker should cite, not Textract's raw
    `Page` field. Confirmed-in-production gap this fixes: Textract's
    `Page` counts from the start of the PDF FILE, with no awareness that
    a manual's own numbering may start later (unnumbered cover/TOC
    pages push everything after them out of sync). On the real manual
    this was built against, a section whose own printed header reads
    "28" sits on the PDF's 29th physical page -- citing raw file-order
    would have told a user "page 29" for content the printed book itself
    labels "28," off by exactly the front-matter page count. Falls back
    to the raw Textract page index for any page where no LAYOUT_PAGE_NUMBER
    block was detected, or where its text isn't cleanly a number (e.g. a
    roman-numeral front-matter page, or a bad OCR read) -- a slightly
    file-order-relative citation beats one built from garbage text.
    """
    blocks_by_id = {b["Id"]: b for b in blocks if b.get("Id")}

    printed_page_labels = {}
    for b in blocks:
        if b.get("BlockType") != "LAYOUT_PAGE_NUMBER":
            continue
        raw_page = b.get("Page", 1)
        if raw_page in printed_page_labels:
            continue  # first one found for this raw page wins; defensive against duplicates
        label_text = _resolve_block_text(b, blocks_by_id, join=" ").strip()
        if label_text.isdigit():
            printed_page_labels[raw_page] = int(label_text)

    def _cited_page(raw_page):
        return printed_page_labels.get(raw_page, raw_page)

    layout_blocks = [b for b in blocks if str(b.get("BlockType", "")).startswith("LAYOUT_")]
    layout_blocks = [b for b in layout_blocks if b.get("BlockType") not in _LAYOUT_SKIP_TYPES]

    def _sort_key(b):
        page = b.get("Page", 1)
        geom = (b.get("Geometry") or {}).get("BoundingBox") or {}
        return (page, geom.get("Top", 0.0), geom.get("Left", 0.0))

    layout_blocks.sort(key=_sort_key)

    # Marker placement is NOT simply "on every page change" -- confirmed
    # by direct testing against MarkdownHeaderTextSplitter that a marker
    # placed immediately BEFORE a header line gets swept into the
    # PRECEDING section's content (the splitter treats the header line
    # itself as the exact boundary, so anything sitting just before it,
    # including a marker meant for the section that's ABOUT to start,
    # ends up attributed to the section that's ending). The fix: emit a
    # header's own page marker AFTER the header line, as the first line
    # of that section's body, unconditionally (not just on a page
    # change) -- every section MarkdownHeaderTextSplitter produces then
    # reliably starts with its own accurate marker regardless of where
    # the split boundary falls. Non-header content blocks keep the
    # simpler "mark only when the page actually changes" behavior, since
    # they don't sit at a boundary the splitter will cut on.
    lines = []
    last_marked_page = None
    for block in layout_blocks:
        page = block.get("Page", 1)
        block_type = block.get("BlockType")

        if block_type in _LAYOUT_HEADER_TYPES:
            text = _resolve_block_text(block, blocks_by_id, join=" ")
            if text:
                lines.append(f"{_LAYOUT_HEADER_TYPES[block_type]} {text}")
                lines.append(f"<!-- page:{_cited_page(page)} -->")
                last_marked_page = page  # raw page, for change-detection below -- see _cited_page()
            lines.append("")
            continue

        if page != last_marked_page:
            lines.append(f"<!-- page:{_cited_page(page)} -->")
            last_marked_page = page

        if block_type == "LAYOUT_LIST":
            # Each child of a LAYOUT_LIST is one list item's LINE -- kept
            # as separate bullets rather than flattened into one
            # paragraph, matching the visual structure Textract detected.
            for rel in block.get("Relationships", []) or []:
                if rel.get("Type") != "CHILD":
                    continue
                for child_id in rel.get("Ids", []):
                    child = blocks_by_id.get(child_id)
                    if child and child.get("Text"):
                        lines.append(f"- {child['Text']}")

        else:
            # LAYOUT_TEXT, LAYOUT_TABLE, LAYOUT_KEY_VALUE, and anything
            # else not explicitly handled above -- plain paragraph text.
            text = _resolve_block_text(block, blocks_by_id, join="\n")
            if text:
                lines.append(text)

        lines.append("")  # blank line between blocks -- standard Markdown paragraph separation

    return "\n".join(lines)


def extract_markdown_from_pdf(file_bytes: bytes, file_url: str) -> str:
    """
    PDF entry point for the Markdown + page-marker extraction path (see
    _blocks_to_markdown()) -- what _run_ingest() actually calls for PDFs
    now. extract_text_from_pdf()/_extract_pdf_page_texts() below are kept
    unchanged, still used by extract_images_from_pdf() for page-level
    image context_text (out of scope for this change) and anything else
    that specifically wants flat text with no structure.
    """
    try:
        blocks = _textract_analyze_layout(file_bytes, file_url)
        return _blocks_to_markdown(blocks)
    except Exception as e:
        log.error(f"PDF layout/Markdown extraction error: {e}")
        raise


def _extract_pdf_page_texts(file_bytes: bytes, file_url: str) -> list:
    """
    TEXTRACT-ONLY BUILD: every PDF is sent through Textract, every time --
    no pypdf native-text pass, no scanned-vs-clean check, no threshold.
    Returns a list of per-page text, one entry per page, in page order.

    pypdf is still used here, but only to get the page COUNT (so the
    returned list has the right length/order even for a page Textract's
    result happens not to cover) -- unlike the hybrid version, its
    extract_text() is never called, so there's no free native-text pass
    being wasted or compared against.

    No fallback on a Textract failure: the hybrid version can fall back
    to the native text it already extracted if Textract errors or times
    out, but this build never extracts native text in the first place,
    so a Textract error here propagates up and fails the whole ingest
    job (see extract_text_from_pdf's except-clause). That's a real
    behavioral difference worth knowing about when comparing the two --
    this build has strictly more exposure to Textract-side failures
    (job failure, timeout, throttling, IAM/network issues) than the
    hybrid one, because it depends on Textract for 100% of documents
    instead of only the ones that actually need OCR.

    Shared by extract_text_from_pdf() (joins the non-empty pages into
    the document's full text) and extract_images_from_pdf() (uses each
    page's text as that image's context_text).
    """
    if not TEXTRACT_ENABLED:
        raise RuntimeError(
            "TEXTRACT_ENABLED is False -- this Textract-only build has no "
            "extraction path at all without it (unlike the hybrid version, "
            "which still has native pypdf extraction to fall back on)."
        )

    import pypdf
    reader     = pypdf.PdfReader(io.BytesIO(file_bytes))
    num_pages  = len(reader.pages)

    log.info(
        f"Sending {num_pages}-page PDF to Textract unconditionally "
        f"(Textract-only build -- no native-text check performed)."
    )
    textract_pages = _textract_extract_document(file_bytes, file_url)

    page_texts = [textract_pages.get(i + 1, "").strip() for i in range(num_pages)]
    filled = sum(1 for t in page_texts if t)
    log.info(f"Textract returned text for {filled}/{num_pages} page(s).")
    return page_texts


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes, file_url: str) -> str:
    try:
        page_texts = _extract_pdf_page_texts(file_bytes, file_url)
        return "\n\n".join(t for t in page_texts if t)
    except Exception as e:
        log.error(f"PDF extraction error: {e}")
        raise


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Unchanged behavior for any existing caller that wants flat text.
    extract_markdown_from_docx() below is the structure-aware version
    _run_ingest() actually uses now.
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        log.error(f"DOCX extraction error: {e}")
        raise


_DOCX_HEADING_STYLE_RE = re.compile(r"^Heading\s+(\d)$", re.IGNORECASE)


def extract_markdown_from_docx(file_bytes: bytes) -> str:
    """
    Same principle as the PDF path switching to Textract's LAYOUT blocks:
    detect structure from the source document's OWN formatting, not a
    guess at a text pattern. Word's built-in paragraph styles
    ("Heading 1", "Heading 2", ..., "Title") are a real, reliable
    structural signal python-docx exposes directly via
    paragraph.style.name -- no OCR, no regex-on-casing guess. A DOCX
    manual authored with Word's built-in heading styles gets real
    section-aware chunking through the same chunk_markdown() as PDFs; one
    that doesn't use heading styles at all degrades to every paragraph
    being unheaded plain text (chunk_markdown() then falls back to one
    big unsectioned chunk sequence via its size-bounded sub-split) -- not
    worse than before this change, just not improved for that document.

    DOCX has no fixed "page" concept at the paragraph-extraction level
    (Word reflows; page breaks depend on rendering, not source content),
    so no `<!-- page:N -->` markers are emitted here -- DOCX-sourced
    chunks legitimately have page_start/page_end = NULL, same as the
    migration comment documents.
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style_name = (p.style.name if p.style else "") or ""
            m = _DOCX_HEADING_STYLE_RE.match(style_name.strip())
            if m:
                level = min(int(m.group(1)), 4)
                lines.append(f"{'#' * level} {text}")
            elif style_name.strip().lower() == "title":
                lines.append(f"# {text}")
            else:
                lines.append(text)
        return "\n\n".join(lines)
    except Exception as e:
        log.error(f"DOCX markdown extraction error: {e}")
        raise


def extract_text(file_bytes: bytes, filename: str, file_url: str) -> str:
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes, file_url)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def extract_markdown(file_bytes: bytes, filename: str, file_url: str) -> str:
    """
    Structure-aware entry point -- what _run_ingest() calls now instead
    of extract_text(). Mirrors extract_text()'s dispatch exactly, just
    routing to the Markdown-producing extractor for each file type.
    """
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_markdown_from_pdf(file_bytes, file_url)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        return extract_markdown_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


# ── Image extraction (PDF only for now) ──────────────────────────────────────

def extract_images_from_pdf(file_bytes: bytes, file_url: str) -> list:
    """
    Extract embedded images from a PDF, paired with the extracted text
    of the page they appear on (used as coarse "context_text" for later
    keyword matching against a component name).

    Returns a list of dicts: {page_number, image_bytes, extension, context_text}
    Skips images smaller than MIN_IMAGE_BYTES (logos/icons/watermarks).

    Page text now comes from _extract_pdf_page_texts(), the same
    native-or-Textract extraction extract_text_from_pdf() uses, so a
    scanned page's images still get real context_text instead of "" --
    note this means the PDF gets parsed twice (once here via pypdf for
    the image objects themselves, once inside _extract_pdf_page_texts
    for text/Textract). That redundancy is harmless correctness-wise and
    not worth optimizing away while this whole function is gated behind
    IMAGE_EXTRACTION_ENABLED = False; revisit if/when it's turned back on.

    NOTE: DOCX image extraction is not implemented yet -- extract_text_from_docx
    only pulls paragraph text. If DOCX manuals with embedded images become
    common, this needs a matching extract_images_from_docx() using
    doc.part.rels to walk embedded image parts.
    """
    import pypdf

    reader     = pypdf.PdfReader(io.BytesIO(file_bytes))
    page_texts = _extract_pdf_page_texts(file_bytes, file_url)
    images     = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page_texts[page_number - 1] if page_number - 1 < len(page_texts) else ""

        try:
            page_images = list(page.images)
        except Exception as e:
            log.warning(f"Failed to read images on page {page_number}: {e}")
            continue

        for img in page_images:
            try:
                image_bytes = img.data
            except Exception as e:
                log.warning(f"Failed to read image data on page {page_number}: {e}")
                continue

            if len(image_bytes) < MIN_IMAGE_BYTES:
                continue  # likely a logo/icon/watermark, not a component reference

            _, ext = os.path.splitext(getattr(img, "name", "") or "")
            ext = ext.lower() if ext else ".png"

            images.append({
                "page_number":  page_number,
                "image_bytes":  image_bytes,
                "extension":    ext,
                "context_text": page_text.strip(),
            })

    log.info(f"Extracted {len(images)} candidate images (>{MIN_IMAGE_BYTES} bytes) from PDF.")
    return images


# ── BOM item extraction (NEW) ──────────────────────────────────────────────────
#
# Many OEM electronic parts manuals (this one included) lay out components
# as fixed-format rows:
#     "2 0001 341047 02 RAIL-TOP-RH 1.000 EA"
#      Lv Item PartNo  Rev Description         Qty  Un
# The Item column (e.g. "0001") is a small integer that, in a meaningful
# fraction of these manuals, also matches the numbered balloon callouts on
# the facing exploded-view drawing -- but that's NOT being relied on here.
# All this does is opportunistically capture (item_no, part_no) pairs
# wherever this row shape appears in the extracted text, so they're not
# lost. It does nothing else: no image is linked, no balloon is detected,
# no assumption is made about what the item number means for any given
# manual. For manuals that don't use this table format, extract_bom_items
# simply returns an empty list -- safe to call unconditionally. On
# Textract'd (scanned) pages, matching also depends on Textract's read
# quality on that scan -- a bad match is a safe no-op, not an error.

_BOM_LINE_RE = re.compile(
    r"^\d+\s+(?P<item>\d{4})\s+(?P<part>[\w./\"-]+)\s+(?P<rest>.+)$"
)
_BOM_TRAILING_QTY_RE = re.compile(
    r"^(?P<desc>.*?)\s+(?P<qty>[\d.]+)\s+(?P<unit>EA|IN|FT)$"
)


def extract_bom_items(text: str) -> list:
    """
    Best-effort, regex-only extraction of (item_no, part_no, description)
    triples from OEM-style parts-list rows in already-extracted page text.
    No OCR, no model call -- just string matching over text we already have.

    Returns a list of dicts: {item_no, part_no, description}. Empty list
    if the text doesn't contain rows in this shape (the common case for
    most manuals) -- callers don't need to check the manual's format
    before calling this.
    """
    items = []
    for line in text.splitlines():
        line = line.strip()
        m = _BOM_LINE_RE.match(line)
        if not m:
            continue

        item_no = m.group("item")
        if item_no == "0000":
            continue  # assembly header / kit-reference row, not a numbered part

        part_no = m.group("part")
        rest    = m.group("rest").strip()

        m2 = _BOM_TRAILING_QTY_RE.match(rest)
        description = m2.group("desc").strip() if m2 else rest

        # Strip a leading 2-digit revision code before the real description,
        # e.g. "02 RAIL-TOP-RH" -> "RAIL-TOP-RH". Cosmetic only -- this
        # field isn't persisted anywhere yet (only item_no/part_no are
        # encoded into the chunk prefix), but keep it clean for whenever
        # it is.
        m3 = re.match(r"^\d{2}\s+(.*)$", description)
        if m3:
            description = m3.group(1)

        items.append({
            "item_no":     item_no,
            "part_no":     part_no,
            "description": description[:120],  # defensive cap
        })

    return items


def _encode_bom_items(bom_items: list) -> str:
    """
    Compact "item=part,item=part" encoding for embedding in the chunk's
    content prefix. Description is intentionally omitted here -- it's
    already present in the chunk's own text for embedding/full-text
    search, so repeating it in the prefix would just be redundant bytes.
    """
    return ",".join(f"{b['item_no']}={b['part_no']}" for b in bom_items)


# ── Chunking: Markdown structural + page-range tagging ────────────────────────
#
# SUPERSEDES two prior versions of this section, in order:
#   1. Pure CHUNK_SIZE-word sliding-window chunking over the entire
#      document as one flattened string -- no awareness of page breaks or
#      section structure at all. Confirmed-in-production bug: a
#      "SEQUENCE OF OPERATION" section got sliced across windows diluted
#      by unrelated neighboring text, so NOT ONE resulting chunk started
#      with or represented that section -- semantic_search() for "what is
#      the sequence of operation" scored under 0.20 against every chunk
#      of that document, and the assistant told a user the content wasn't
#      documented even though it was sitting in the manual.
#   2. An ALL-CAPS-regex header detector, splitting on section headers
#      before windowing. Fixed (1) for that one manual, but was reverse-
#      engineered from that manual's specific casing convention -- caught
#      (correctly) as not generalizing to a different manufacturer's
#      Title Case or numbered headers, since it was never verified
#      against more than one document.
#
# This version fixes the generalization problem: it splits on REAL
# document structure (Markdown headings that extract_markdown_from_pdf()/
# extract_markdown_from_docx() produce from Textract's LAYOUT-classified
# blocks or Word's own heading styles -- see those functions) via
# langchain_text_splitters.MarkdownHeaderTextSplitter, instead of guessing
# at any one manual's formatting convention. A section this splitter
# produces can still be long (e.g. a full "TROUBLE SHOOTING TIPS"
# chapter), so the same CHUNK_SIZE/CHUNK_OVERLAP word-windowing as before
# still runs WITHIN each header-bounded section -- it just never crosses a
# section boundary anymore, and every resulting chunk is prefixed with its
# full heading path (e.g. "TITLE > SEQUENCE OF OPERATION"), not just the
# immediate header, reconstructed from the splitter's own metadata rather
# than re-guessed.
#
# ALSO NEW: extracts each chunk's page_start/page_end from the
# `<!-- page:N -->` markers embedded in the Markdown (see
# _blocks_to_markdown()), then strips them from the text before it's
# embedded or shown to anyone -- bookkeeping, like the equipment_id/doc_id
# content prefix, not content. save_chunks() persists these as real
# columns (migration_003_knowledge_embeddings_page_range.sql) so
# retrieval.py can surface a real "Found on Page 243" citation.
#
# NOTE: this only changes chunking for documents ingested (or
# re-ingested) AFTER this ships -- it does nothing for chunks already
# sitting in knowledge_embeddings from a prior ingestion. Those keep
# whatever the chunker in place at ingest time produced (flat-window text,
# no page range) until that document is explicitly re-ingested.
_MARKDOWN_HEADERS = [("#", "h1"), ("##", "h2"), ("###", "h3"), ("####", "h4")]


def _extract_and_strip_page_range(text: str):
    """
    Returns (clean_text, page_start, page_end) -- page_start/page_end are
    None if `text` contains no `<!-- page:N -->` markers at all (e.g. a
    DOCX source, which never has any, or a chunk that somehow ended up
    with none) -- treated as "unknown," never defaulted to a guessed page.
    """
    pages = [int(m) for m in _PAGE_MARKER_RE.findall(text)]
    clean = _PAGE_MARKER_RE.sub("", text)
    clean = re.sub(r"\n{3,}", "\n\n", clean).strip()
    if not pages:
        return clean, None, None
    return clean, min(pages), max(pages)


def chunk_markdown(markdown_text: str) -> list:
    """
    Returns a list of {"text": str, "page_start": int|None, "page_end": int|None}
    dicts, ready for save_chunks() -- this is what _run_ingest() calls
    now instead of chunk_text(). See the module comment above for the
    two-pass design: structural split first (MarkdownHeaderTextSplitter),
    then a size-bounded sub-split within each resulting section.
    """
    from langchain_text_splitters import MarkdownHeaderTextSplitter

    splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on=_MARKDOWN_HEADERS, strip_headers=False,
    )
    sections = splitter.split_text(markdown_text)

    results = []
    for section in sections:
        heading_path = " > ".join(
            section.metadata[k] for k in ("h1", "h2", "h3", "h4")
            if k in section.metadata
        )

        words = section.page_content.split()
        if not words:
            continue

        start = 0
        while start < len(words):
            end         = min(start + CHUNK_SIZE, len(words))
            window_text = " ".join(words[start:end])
            clean_text, page_start, page_end = _extract_and_strip_page_range(window_text)

            if clean_text:
                full_text = f"{heading_path}\n{clean_text}" if heading_path else clean_text
                results.append({
                    "text": full_text,
                    "page_start": page_start,
                    "page_end": page_end,
                })

            if end == len(words):
                break
            start = end - CHUNK_OVERLAP

    return results


# ── DB helpers: text chunks ───────────────────────────────────────────────────

def clear_document_chunks(conn, file_url: str, company_id: str):
    """
    Remove all chunks for a specific document identified by file URL.
    Matches on the doc_id prefix embedded in the content column.
    Called when a document is deleted or re-uploaded.
    Does not affect other documents on the same node.
    """
    doc_uuid = url_to_uuid(file_url)
    with conn.cursor() as cur:
        cur.execute("""
            DELETE FROM knowledge_embeddings
            WHERE company_id = %s::uuid
            AND source_type = 'manual'
            AND content LIKE %s
        """, (company_id, f"%doc_id:{doc_uuid}%"))
        deleted = cur.rowcount
    conn.commit()
    log.info(f"Removed {deleted} chunks for document {file_url}")
    return deleted


def clear_node_chunks(conn, equipment_id: str, company_id: str):
    """
    Remove ALL manual chunks for a node (equipment or component).
    Called when a node is deleted. Cleans up all its documents at once.

    Filters on the real equipment_id column (see
    migration_002_knowledge_embeddings_equipment_id.sql) rather than the
    old content LIKE '%equipment_id:{id}%' text match -- exact and
    indexed instead of a substring scan. Requires that migration to have
    run AND its backfill to be complete, or pre-migration rows (not yet
    backfilled) won't be matched here even though they'd have matched
    the old LIKE-based query.
    """
    with conn.cursor() as cur:
        cur.execute("""
            DELETE FROM knowledge_embeddings
            WHERE source_type = 'manual'
            AND company_id = %s::uuid
            AND equipment_id = %s::uuid
        """, (company_id, equipment_id))
        deleted = cur.rowcount
    conn.commit()
    log.info(f"Removed {deleted} chunks for node {equipment_id}")
    return deleted


# ── DB + S3 helpers: manual images ────────────────────────────────────────────

def _s3_client():
    return boto3.client("s3", region_name=AWS_REGION)


def clear_document_images(conn, file_url: str, company_id: str):
    """
    Remove all manual images for a specific document: deletes both the
    S3 objects and the DB rows. Mirrors clear_document_chunks() so a
    re-uploaded document doesn't accumulate duplicate images over time.
    """
    doc_uuid = url_to_uuid(file_url)
    with conn.cursor() as cur:
        cur.execute("""
            SELECT s3_key FROM equipment_manual_images
            WHERE company_id = %s::uuid AND doc_id = %s::uuid
        """, (company_id, doc_uuid))
        keys = [row["s3_key"] for row in cur.fetchall()]

    if keys:
        s3 = _s3_client()
        # S3 delete_objects takes up to 1000 keys per call
        for i in range(0, len(keys), 1000):
            batch = keys[i:i + 1000]
            s3.delete_objects(
                Bucket=S3_BUCKET,
                Delete={"Objects": [{"Key": k} for k in batch]},
            )

    with conn.cursor() as cur:
        cur.execute("""
            DELETE FROM equipment_manual_images
            WHERE company_id = %s::uuid AND doc_id = %s::uuid
        """, (company_id, doc_uuid))
        deleted = cur.rowcount
    conn.commit()
    log.info(f"Removed {deleted} manual images (and their S3 objects) for document {file_url}")
    return deleted


def clear_node_images(conn, equipment_id: str, company_id: str):
    """
    Remove ALL manual images for a node: S3 objects + DB rows.
    Mirrors clear_node_chunks(). Called when a node is deleted.
    """
    with conn.cursor() as cur:
        cur.execute("""
            SELECT s3_key FROM equipment_manual_images
            WHERE company_id = %s::uuid AND equipment_id = %s::uuid
        """, (company_id, equipment_id))
        keys = [row["s3_key"] for row in cur.fetchall()]

    if keys:
        s3 = _s3_client()
        for i in range(0, len(keys), 1000):
            batch = keys[i:i + 1000]
            s3.delete_objects(
                Bucket=S3_BUCKET,
                Delete={"Objects": [{"Key": k} for k in batch]},
            )

    with conn.cursor() as cur:
        cur.execute("""
            DELETE FROM equipment_manual_images
            WHERE company_id = %s::uuid AND equipment_id = %s::uuid
        """, (company_id, equipment_id))
        deleted = cur.rowcount
    conn.commit()
    log.info(f"Removed {deleted} manual images (and their S3 objects) for node {equipment_id}")
    return deleted


def save_manual_images(conn, images: list, file_url: str, equipment_id: str, company_id: str):
    """
    Upload extracted images to S3 (privately -- key stored, not a
    public URL) and record one row per image in equipment_manual_images
    for later keyword lookup during PM strategy generation.
    """
    if not images:
        return

    doc_uuid = url_to_uuid(file_url)
    s3 = _s3_client()

    rows = []
    for idx, img in enumerate(images):
        ext = img["extension"] if img["extension"] in _EXTENSION_CONTENT_TYPES else ".png"
        content_type = _EXTENSION_CONTENT_TYPES.get(ext, "application/octet-stream")

        s3_key = f"manual-images/{company_id}/{equipment_id}/{doc_uuid}/{img['page_number']}_{idx}{ext}"

        try:
            s3.put_object(
                Bucket=S3_BUCKET,
                Key=s3_key,
                Body=img["image_bytes"],
                ContentType=content_type,
            )
        except Exception as e:
            log.warning(f"Failed to upload manual image {s3_key} to S3: {e}")
            continue

        rows.append((
            str(uuid.uuid4()),
            company_id,
            equipment_id,
            doc_uuid,
            s3_key,
            img["page_number"],
            img["context_text"][:5000],  # cap context text length defensively
        ))

    if not rows:
        return

    sql = """
        INSERT INTO equipment_manual_images
            (id, company_id, equipment_id, doc_id, s3_key, page_number,
             context_text, created_at, updated_at)
        VALUES %s
    """
    psycopg2.extras.execute_values(
        conn.cursor(), sql, rows,
        template="(%s, %s::uuid, %s::uuid, %s::uuid, %s, %s, %s, NOW(), NOW())",
        page_size=100,
    )
    conn.commit()
    log.info(f"Saved {len(rows)} manual images for document {file_url}")


# ── Concurrent embedding ───────────────────────────────────────────────────────

async def _embed_one(executor_loop, semaphore, index: int, content: str):
    """
    Embed a single chunk, bounded by a semaphore so we never have more
    than EMBED_CONCURRENCY calls in flight at once. Runs the blocking
    get_embedding() call in a thread executor so it doesn't block the loop.
    """
    async with semaphore:
        embedding = await executor_loop.run_in_executor(None, get_embedding, content)
        return index, embedding


async def embed_chunks_concurrently(contents: list) -> list:
    """
    Embed all chunk contents concurrently, bounded by EMBED_CONCURRENCY.
    Returns embeddings in the same order as the input contents list.
    """
    loop      = asyncio.get_event_loop()
    semaphore = asyncio.Semaphore(EMBED_CONCURRENCY)

    tasks = [
        _embed_one(loop, semaphore, i, content)
        for i, content in enumerate(contents)
    ]

    results = await asyncio.gather(*tasks)
    results.sort(key=lambda r: r[0])
    return [embedding for _, embedding in results]


def save_chunks(conn, chunks: list, file_url: str, equipment_id: str, company_id: str):
    """
    Embed and save all chunks. Each chunk gets:
      - a unique source_id derived from (doc_uuid, chunk_index) via uuid5
        so re-ingesting the same file produces the same chunk UUIDs
      - a real knowledge_embeddings.equipment_id column value (see
        migration_002_knowledge_embeddings_equipment_id.sql -- THAT
        MIGRATION MUST BE APPLIED before this function will work; the
        INSERT below will fail with "column equipment_id does not exist"
        otherwise) -- this is what retrieval.py's semantic_search() now
        filters on directly for equipment-scoped chat retrieval
      - a content prefix with equipment_id, doc_id, and (NEW) any BOM
        item/part-number pairs found in that chunk's text -- still
        written for document-scoped deletion (clear_document_chunks,
        doc_id has no column of its own) and for
        fetch_all_manual_chunks() in generate_pm_strategy.py, which
        parses this fixed-shape prefix. The equipment_id copy in this
        prefix is now redundant with the new column for retrieval
        purposes, but kept so that parser doesn't break.
      - real page_start/page_end columns (migration_003_knowledge_
        embeddings_page_range.sql -- MUST be applied before this INSERT
        will work, same ordering caution as migration_002) -- NOT folded
        into the content prefix, since that convention is already
        flagged for retirement (see SKILL.md) and there's no reason to
        grow it further.

    The content prefix is always exactly three "key:value | " segments
    followed by the chunk text -- "bom_items" is included even when empty
    so the prefix has a fixed, predictable shape for parsing later
    (see fetch_all_manual_chunks in generate_pm_strategy.py).

    `chunks` is now a list of {"text", "page_start", "page_end"} dicts
    (see chunk_markdown() in ingest_document.py), not a list of plain
    strings -- the shape changed together with the Markdown structural
    chunking rewrite.

    Embeddings are generated concurrently (bounded by EMBED_CONCURRENCY)
    instead of one at a time, so large documents (50-100+ pages) complete
    well within the Lambda timeout window.
    """
    doc_uuid = url_to_uuid(file_url)

    contents    = []
    page_ranges = []
    chunks_with_bom = 0
    for chunk in chunks:
        chunk_text_value = chunk["text"]
        bom_items = extract_bom_items(chunk_text_value)
        if bom_items:
            chunks_with_bom += 1
        bom_field = _encode_bom_items(bom_items)
        contents.append(
            f"equipment_id:{equipment_id} | doc_id:{doc_uuid} | bom_items:{bom_field} | {chunk_text_value}"
        )
        page_ranges.append((chunk.get("page_start"), chunk.get("page_end")))

    if chunks_with_bom:
        log.info(f"Found parts-list rows in {chunks_with_bom}/{len(chunks)} chunks.")

    log.info(f"Embedding {len(contents)} chunks with concurrency={EMBED_CONCURRENCY}...")
    embeddings = asyncio.run(embed_chunks_concurrently(contents))
    log.info(f"Finished embedding {len(embeddings)} chunks.")

    rows = []
    for i, (content, embedding) in enumerate(zip(contents, embeddings)):
        emb_str   = "[" + ",".join(map(str, embedding)) + "]"
        record_id = str(uuid.uuid4())
        chunk_source_id = str(uuid.uuid5(uuid.UUID(doc_uuid), str(i)))
        page_start, page_end = page_ranges[i]

        rows.append((
            record_id,
            company_id,
            "manual",
            chunk_source_id,
            equipment_id,
            content,
            emb_str,
            page_start,
            page_end,
        ))

    sql = """
        INSERT INTO knowledge_embeddings
            (id, company_id, source_type, source_id, equipment_id, content,
             embedding, page_start, page_end, created_at, updated_at)
        VALUES %s
    """
    psycopg2.extras.execute_values(
        conn.cursor(), sql, rows,
        template="(%s, %s::uuid, %s, %s::uuid, %s::uuid, %s, %s::vector, %s, %s, NOW(), NOW())",
        page_size=100,
    )
    conn.commit()
    log.info(f"Saved {len(rows)} chunks for document {file_url}")


# ── Synchronous core (used by both sync and async entry points) ───────────────

def _run_ingest(file_url: str, equipment_id: str, company_id: str) -> dict:
    """
    The actual ingest work. Shared by the synchronous ingest() function
    and the async run_ingest_job() background runner.
    """
    log.info(f"Ingesting document for equipment {equipment_id}: {file_url}")

    filename   = file_url.split("/")[-1].split("?")[0]
    response   = requests.get(file_url, timeout=30)
    response.raise_for_status()
    file_bytes = response.content
    log.info(f"Downloaded {len(file_bytes)} bytes from {filename}")

    # extract_markdown() (Textract LAYOUT for PDF, Word heading styles for
    # DOCX) replaces extract_text() here -- see that function and
    # chunk_markdown() for why. extract_text()/chunk_text() are kept in
    # the file, unused by this path, in case something else still wants
    # flat unstructured text.
    markdown_text = extract_markdown(file_bytes, filename, file_url)
    if not markdown_text.strip():
        if filename.lower().endswith(".pdf"):
            raise ValueError(
                "No text could be extracted from the document. This build "
                "sends every PDF through Textract (LAYOUT analysis) "
                "unconditionally, so an empty result here means Textract "
                "itself returned nothing for every page -- check the "
                "CloudWatch logs above for a Textract error: job failure, "
                "timeout, TEXTRACT_ENABLED is False, or the Lambda "
                "execution role missing textract:StartDocumentAnalysis / "
                "textract:GetDocumentAnalysis / s3:GetObject. (Layout "
                "analysis needs different IAM actions than the old plain "
                "text-detection path -- StartDocumentAnalysis/"
                "GetDocumentAnalysis, not StartDocumentTextDetection/"
                "GetDocumentTextDetection -- if this used to work and now "
                "fails with an access-denied error, that permission is "
                "the first thing to check.)"
            )
        raise ValueError("No text could be extracted from the document.")
    log.info(f"Extracted {len(markdown_text.split())} words (Markdown).")

    chunks = chunk_markdown(markdown_text)
    log.info(f"Created {len(chunks)} chunks.")
    with_pages = sum(1 for c in chunks if c.get("page_start") is not None)
    log.info(f"{with_pages}/{len(chunks)} chunks have a page range attached.")

    # Image extraction is PDF-only, and gated behind IMAGE_EXTRACTION_ENABLED
    # (currently False -- see module docstring). Skipping this entirely
    # while disabled is the whole point: no per-page image pulls, no
    # per-image S3 uploads, so ingest is just download -> text -> chunk ->
    # embed. See extract_images_from_pdf's docstring for the DOCX gap
    # whenever this is turned back on.
    images = []
    if IMAGE_EXTRACTION_ENABLED and filename.lower().endswith(".pdf"):
        try:
            images = extract_images_from_pdf(file_bytes, file_url)
        except Exception as e:
            # Image extraction failing should never block text ingestion --
            # log and continue with an empty image set.
            log.error(f"Image extraction failed for {filename}, continuing without images: {e}")
            images = []

    conn = get_db_connection()
    try:
        clear_document_chunks(conn, file_url, company_id)
        save_chunks(conn, chunks, file_url, equipment_id, company_id)

        if IMAGE_EXTRACTION_ENABLED:
            clear_document_images(conn, file_url, company_id)
            save_manual_images(conn, images, file_url, equipment_id, company_id)
    finally:
        conn.close()

    return {
        "equipment_id": equipment_id,
        "company_id":   company_id,
        "source_type":  "manual",
        "filename":     filename,
        "words":        len(markdown_text.split()),
        "chunks":       len(chunks),
        "images":       len(images),
        "status":       "ingested"
    }


# ── Main functions ────────────────────────────────────────────────────────────

def ingest(file_url: str, equipment_id: str, company_id: str) -> dict:
    """
    Synchronous ingest, kept for any direct/internal callers.
    Most traffic should go through run_ingest_job via SQS instead.
    """
    return _run_ingest(file_url, equipment_id, company_id)


def run_ingest_job(job_id: str, file_url: str, equipment_id: str, company_id: str):
    """
    Called when Lambda is triggered by SQS for a document_ingest job.
    Runs the ingest, then updates the jobs table with the result.
    """
    try:
        result = _run_ingest(file_url, equipment_id, company_id)

        conn = get_db_connection()
        try:
            with conn.cursor() as cur:
                cur.execute("""
                    UPDATE pm_strategy_jobs
                    SET status = 'ready', result = %s::jsonb, updated_at = NOW()
                    WHERE id = %s::uuid
                """, (psycopg2.extras.Json(result), job_id))
            conn.commit()
        finally:
            conn.close()

        log.info(f"DOCUMENT INGEST JOB COMPLETE [{job_id}]")

    except Exception as e:
        log.error(f"DOCUMENT INGEST JOB ERROR [{job_id}]: {str(e)}")
        conn = get_db_connection()
        try:
            with conn.cursor() as cur:
                cur.execute("""
                    UPDATE pm_strategy_jobs
                    SET status = 'failed', error = %s, updated_at = NOW()
                    WHERE id = %s::uuid
                """, (str(e), job_id))
            conn.commit()
        finally:
            conn.close()


def delete_document(file_url: str, company_id: str) -> dict:
    """
    Remove all chunks AND images for a specific document.
    Called by Benjamin when a document is deleted from a node.
    """
    conn = get_db_connection()
    try:
        deleted_chunks = clear_document_chunks(conn, file_url, company_id)
        deleted_images = clear_document_images(conn, file_url, company_id)
    finally:
        conn.close()
    return {
        "file_url":        file_url,
        "chunks_deleted":  deleted_chunks,
        "images_deleted":  deleted_images,
        "status":          "deleted",
    }


def delete_node(equipment_id: str, company_id: str) -> dict:
    """
    Remove all manual chunks AND images for a node.
    Called by Benjamin when an equipment or component node is deleted.
    """
    conn = get_db_connection()
    try:
        deleted_chunks = clear_node_chunks(conn, equipment_id, company_id)
        deleted_images = clear_node_images(conn, equipment_id, company_id)
    finally:
        conn.close()
    return {
        "equipment_id":    equipment_id,
        "chunks_deleted":  deleted_chunks,
        "images_deleted":  deleted_images,
        "status":          "deleted",
    }